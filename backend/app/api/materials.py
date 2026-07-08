"""
教学资料与题库管理 API — PDF上传、AI出题、作业发布。
"""

from __future__ import annotations

import json
import os
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any

from fastapi import APIRouter, Form, HTTPException, UploadFile
from app.core.llm import chat_json
from app.models.schemas import APIResponse

router = APIRouter(prefix="/api/materials", tags=["教学资料"])

# 存储路径
MATERIALS_DIR = Path(__file__).parent.parent.parent / "knowledge_base" / "materials"
QUESTIONS_DIR = Path(__file__).parent.parent.parent / "knowledge_base" / "questions"
MATERIALS_DIR.mkdir(parents=True, exist_ok=True)
QUESTIONS_DIR.mkdir(parents=True, exist_ok=True)

# 内存索引（生产环境应使用数据库）
_materials_index: dict[str, dict] = {}
_questions_index: dict[str, dict] = {}


def _load_indexes():
    """加载持久化索引。"""
    idx_file = MATERIALS_DIR / "_index.json"
    if idx_file.exists():
        try:
            data = json.loads(idx_file.read_text(encoding="utf-8"))
            _materials_index.update(data.get("materials", {}))
            _questions_index.update(data.get("questions", {}))
        except Exception:
            pass


def _save_indexes():
    """保存索引到磁盘。"""
    idx_file = MATERIALS_DIR / "_index.json"
    idx_file.write_text(
        json.dumps({"materials": _materials_index, "questions": _questions_index}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


# 启动时加载
_load_indexes()


# ── 资料管理 ──────────────────────────────────────────────

@router.post("/upload", response_model=APIResponse)
async def upload_material(file: UploadFile, course: str = Form(""), chapter: str = Form("")):
    """上传教学资料（PDF / Word）。"""
    if not file.filename:
        raise HTTPException(status_code=400, detail="请选择文件")

    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ("pdf", "docx", "doc"):
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: .{ext}（支持 PDF、Word）")

    material_id = str(uuid.uuid4())[:8]
    save_path = MATERIALS_DIR / f"{material_id}_{file.filename}"
    content = await file.read()
    save_path.write_bytes(content)

    # 提取文本（用于后续出题）
    text_content = ""
    pages = 0
    try:
        if ext == "pdf":
            from pypdf import PdfReader
            import io
            reader = PdfReader(io.BytesIO(content))
            pages = len(reader.pages)
            text_content = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif ext in ("docx", "doc"):
            try:
                from docx import Document
                import io
                doc = Document(io.BytesIO(content))
                text_content = "\n".join(p.text for p in doc.paragraphs)
                pages = max(1, len(doc.paragraphs) // 40)  # 粗略估算页数
            except Exception:
                raise RuntimeError("无法解析 .doc 文件，请转换为 .docx 格式后重试")
    except Exception as e:
        text_content = f"（文本提取失败: {str(e)}）"

    info = {
        "id": material_id,
        "filename": file.filename,
        "course": course or "未分类",
        "chapter": chapter or "",
        "size": len(content),
        "size_display": f"{len(content) / 1024:.1f} KB" if len(content) < 1024 * 1024 else f"{len(content) / 1024 / 1024:.1f} MB",
        "pages": pages,
        "text_preview": text_content[:500],
        "text_content": text_content,
        "created_at": datetime.now().isoformat()[:19],
    }
    _materials_index[material_id] = info
    _save_indexes()

    return APIResponse(success=True, message=f"上传成功：{file.filename}", data={
        "id": material_id,
        "filename": file.filename,
        "course": info["course"],
        "size_display": info["size_display"],
        "pages": info["pages"],
        "created_at": info["created_at"],
    })


@router.get("/list", response_model=APIResponse)
async def list_materials(course: str = ""):
    """获取教学资料列表。"""
    items = list(_materials_index.values())
    if course:
        items = [i for i in items if i["course"] == course]
    # 按时间倒序
    items.sort(key=lambda x: x.get("created_at", ""), reverse=True)
    # 不返回完整文本内容（仅预览）
    for item in items:
        item.pop("text_content", None)
    return APIResponse(success=True, data={"total": len(items), "items": items})


@router.get("/detail/{material_id}", response_model=APIResponse)
async def get_material(material_id: str):
    """获取单个资料详情。"""
    info = _materials_index.get(material_id)
    if not info:
        raise HTTPException(status_code=404, detail="资料不存在")
    return APIResponse(success=True, data=info)


@router.get("/download/{material_id}")
async def download_material(material_id: str):
    """下载已上传的教学资料文件。"""
    from fastapi.responses import FileResponse
    info = _materials_index.get(material_id)
    if not info:
        raise HTTPException(status_code=404, detail="资料不存在")
    for f in MATERIALS_DIR.iterdir():
        if f.name.startswith(material_id):
            return FileResponse(
                path=str(f),
                filename=info["filename"],
                media_type="application/octet-stream",
            )
    raise HTTPException(status_code=404, detail="文件不存在，可能已被删除")


@router.delete("/delete/{material_id}", response_model=APIResponse)
async def delete_material(material_id: str):
    """删除教学资料。"""
    info = _materials_index.pop(material_id, None)
    if not info:
        raise HTTPException(status_code=404, detail="资料不存在")
    # 删除文件
    for f in MATERIALS_DIR.iterdir():
        if f.name.startswith(material_id):
            f.unlink(missing_ok=True)
    _save_indexes()
    return APIResponse(success=True, message="已删除")


# ── AI 出题 ────────────────────────────────────────────────

QUESTION_SYSTEM_PROMPT = """【当前任务：学科专业试题智能出题与分层命题】
你是一流学科建设专家型教师，根据教材内容生成标准化本科试题。

出题要求：
1. 分层出题：基础题、提高题、综合应用题、前沿创新题
2. 题型包含：选择、判断、简答、计算、案例分析、论述（按需适配）
3. 每题包含：题目、标准答案、分步评分细则、详细解析、易错点分析
4. 每题绑定知识点标签、教学目标、命题依据（教材来源/页码）
5. 试题规避网络原题，具备本科专业高阶考察性
6. 难度分布合理，覆盖不同认知层次（记忆/理解/应用/分析/评价/创造）

末尾添加 AI 生成标识。

输出 JSON 格式：
{
  "questions": [
    {
      "question": "题目内容",
      "type": "选择题|填空题|简答题|计算题|论述题|案例分析",
      "options": ["A. xxx", "B. xxx", "C. xxx", "D. xxx"],
      "answer": "标准答案",
      "difficulty": "基础|提高|综合|前沿创新",
      "knowledge_point": "知识点名称",
      "teaching_objective": "对应教学目标",
      "source": "命题依据（教材章节/页码/文献）",
      "scoring_rubric": "分步评分细则",
      "common_mistakes": "学生常见易错点",
      "explanation": "详细解析及解题思路",
      "cognitive_level": "记忆|理解|应用|分析|评价|创造",
      "estimated_time": 5
    }
  ]
}"""


@router.post("/generate-questions", response_model=APIResponse)
async def generate_questions(data: dict):
    """
    基于教学资料生成练习题。

    请求体：
    {
        "material_id": "xxx",
        "count": 5,
        "difficulty": "中等",
        "types": ["选择题", "填空题", "简答题"]
    }
    """
    material_id = data.get("material_id", "")
    count = min(data.get("count", 5), 20)
    difficulty = data.get("difficulty", "中等")
    types = data.get("types", ["选择题", "填空题", "简答题"])
    question_types = data.get("question_types", None) or types  # 兼容两种参数名

    info = _materials_index.get(material_id)
    if not info:
        raise HTTPException(status_code=404, detail="资料不存在，请先上传PDF")

    text = info.get("text_content", "")
    if not text or len(text.strip()) < 50:
        raise HTTPException(status_code=400, detail="资料文本内容不足，无法出题")

    # 截取前 8000 字符作为上下文
    context = text[:8000]

    user_prompt = f"""请根据以下教材内容，生成 {count} 道{难度级别}的练习题。

教材内容（{info['filename']} - {info['chapter'] or info['course']}）：
{context}

题目类型：{', '.join(question_types)}
难度级别：{difficulty}
题目数量：{count} 题

请确保覆盖不同知识点，难度分布合理。"""

    try:
        result = chat_json(
            messages=[
                {"role": "system", "content": QUESTION_SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.5,
        )
        questions = result.get("questions", [])
    except Exception as e:
        return APIResponse(success=False, message=f"出题失败：{str(e)[:100]}")

    # 保存题目
    batch_id = str(uuid.uuid4())[:8]
    saved_questions = []
    for i, q in enumerate(questions):
        qid = f"{batch_id}_{i}"
        question_item = {
            "id": qid,
            "batch_id": batch_id,
            "material_id": material_id,
            "material_name": info["filename"],
            "course": info["course"],
            "question": q.get("question", ""),
            "type": q.get("type", "简答题"),
            "options": q.get("options", []),
            "answer": q.get("answer", ""),
            "difficulty": q.get("difficulty", "中等"),
            "knowledge_point": q.get("knowledge_point", ""),
            "explanation": q.get("explanation", ""),
            "estimated_time": q.get("estimated_time", 5),
            "status": "draft",  # draft / published
            "created_at": datetime.now().isoformat()[:19],
        }
        _questions_index[qid] = question_item
        saved_questions.append(question_item)

    _save_indexes()

    return APIResponse(
        success=True,
        message=f"成功生成 {len(saved_questions)} 道题目",
        data={
            "batch_id": batch_id,
            "total": len(saved_questions),
            "questions": saved_questions,
        },
    )


@router.get("/questions", response_model=APIResponse)
async def list_questions(material_id: str = "", status: str = "", course: str = ""):
    """获取题目列表。"""
    items = list(_questions_index.values())
    if material_id:
        items = [i for i in items if i["material_id"] == material_id]
    if status:
        items = [i for i in items if i["status"] == status]
    if course:
        items = [i for i in items if i.get("course") == course]
    items.sort(key=lambda x: x.get("created_at", ""), reverse=True)
    return APIResponse(success=True, data={"total": len(items), "items": items})


@router.post("/questions/update", response_model=APIResponse)
async def update_question(data: dict):
    """更新单道题目。"""
    qid = data.get("id", "")
    if qid not in _questions_index:
        raise HTTPException(status_code=404, detail="题目不存在")
    for key in ("question", "answer", "explanation", "difficulty", "type", "options", "knowledge_point"):
        if key in data:
            _questions_index[qid][key] = data[key]
    _save_indexes()
    return APIResponse(success=True, message="已更新")


@router.post("/publish", response_model=APIResponse)
async def publish_questions(data: dict):
    """
    发布题目给学生端。

    请求体：
    {
        "question_ids": ["id1", "id2"],
        "course": "离散数学",
        "title": "第三章课后练习",
        "deadline": "2026-06-20"
    }
    """
    qids = data.get("question_ids", [])
    course = data.get("course", "")
    title = data.get("title", "练习题")
    deadline = data.get("deadline", "")

    published = []
    for qid in qids:
        if qid in _questions_index:
            _questions_index[qid]["status"] = "published"
            published.append(qid)

    if not published:
        raise HTTPException(status_code=400, detail="没有可发布的题目")

    _save_indexes()

    # 构建发布记录
    publish_record = {
        "id": str(uuid.uuid4())[:8],
        "title": title,
        "course": course,
        "deadline": deadline,
        "question_count": len(published),
        "question_ids": published,
        "created_at": datetime.now().isoformat()[:19],
    }

    pub_file = QUESTIONS_DIR / f"publish_{publish_record['id']}.json"
    pub_file.write_text(json.dumps(publish_record, ensure_ascii=False, indent=2), encoding="utf-8")

    return APIResponse(
        success=True,
        message=f"已成功发布 {len(published)} 道题目",
        data=publish_record,
    )


@router.get("/publish/list", response_model=APIResponse)
async def list_published():
    """获取已发布的作业列表。"""
    records = []
    for f in QUESTIONS_DIR.glob("publish_*.json"):
        try:
            records.append(json.loads(f.read_text(encoding="utf-8")))
        except Exception:
            continue
    records.sort(key=lambda x: x.get("created_at", ""), reverse=True)
    return APIResponse(success=True, data={"total": len(records), "items": records})
